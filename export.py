from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io

def parse_report(report_text):
    sections = {}
    current_section = "intro"
    current_content = []
    lines = report_text.split('\n')
    for line in lines:
        if line.startswith('## '):
            if current_content:
                sections[current_section] = '\n'.join(current_content).strip()
            current_section = line[3:].strip()
            current_content = []
        else:
            current_content.append(line)
    if current_content:
        sections[current_section] = '\n'.join(current_content).strip()
    return sections

def generate_pdf(report_text, query):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           rightMargin=inch, leftMargin=inch,
                           topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    cover_title = ParagraphStyle('CoverTitle', parent=styles['Title'],
                                fontSize=28, textColor=colors.HexColor('#4f46e5'),
                                spaceAfter=20, alignment=TA_CENTER)
    cover_sub = ParagraphStyle('CoverSub', parent=styles['Normal'],
                              fontSize=14, textColor=colors.HexColor('#666666'),
                              spaceAfter=10, alignment=TA_CENTER)
    h1_style = ParagraphStyle('H1', parent=styles['Heading1'],
                             fontSize=18, textColor=colors.HexColor('#1a1a2e'),
                             spaceBefore=20, spaceAfter=10,
                             borderPad=5)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'],
                             fontSize=14, textColor=colors.HexColor('#4f46e5'),
                             spaceBefore=15, spaceAfter=8)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                               fontSize=11, leading=16,
                               spaceAfter=8)
    bullet_style = ParagraphStyle('Bullet', parent=styles['Normal'],
                                 fontSize=11, leading=16,
                                 leftIndent=20, spaceAfter=4,
                                 bulletIndent=10)
    story = []
    story.append(Spacer(1, 2*inch))
    story.append(Paragraph("Research Report", cover_title))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(query, cover_sub))
    story.append(Spacer(1, 0.2*inch))
    from datetime import date
    story.append(Paragraph(f"Generated on {date.today().strftime('%B %d, %Y')}", cover_sub))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Powered by Multi-Agent Research Assistant", cover_sub))
    story.append(PageBreak())
    lines = report_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1*inch))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], h1_style))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], h2_style))
        elif line.startswith('* ') or line.startswith('- '):
            clean = line[2:]
            clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', clean)
            story.append(Paragraph(f"• {clean}", bullet_style))
        elif re.match(r'^\d+\.', line):
            clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(clean, bullet_style))
        elif line.startswith('**') and line.endswith('**'):
            story.append(Paragraph(f"<b>{line[2:-2]}</b>", body_style))
        else:
            clean = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(clean, body_style))
    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_docx(report_text, query):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    title = doc.add_heading('Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    subtitle = doc.add_paragraph(query)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    from datetime import date
    date_para = doc.add_paragraph(f"Generated on {date.today().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    lines = report_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=2)
            h.runs[0].font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line, style='List Number')
        else:
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            p = doc.add_paragraph(clean)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer